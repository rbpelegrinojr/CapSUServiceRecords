import io
import os
from datetime import datetime

from flask import Blueprint, send_file, flash, redirect, url_for
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from models.employee import Employee
from models.service_record import ServiceRecord
from models.setting import Setting

print_doc_bp = Blueprint('print_doc', __name__)


def _get_setting(key, default=''):
    s = Setting.query.filter_by(key=key).first()
    return s.value if s else default


def _fmt_salary(val):
    if val is None:
        return ''
    return f'{val:,.2f}'


def _replace_text_in_paragraph(para, replacements):
    """Replace {{key}} tokens in a paragraph, preserving runs."""
    full_text = ''.join(run.text for run in para.runs)
    new_text = full_text
    for key, value in replacements.items():
        new_text = new_text.replace('{{' + key + '}}', value)
    if new_text != full_text and para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ''


def _set_cell_text(cell, text, bold=False, font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    para = cell.paragraphs[0]
    para.alignment = alignment
    run = para.add_run(str(text) if text else '')
    run.font.size = Pt(font_size)
    run.bold = bold


def _set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for border_name in ('top', 'left', 'bottom', 'right'):
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single')
        border_el.set(qn('w:sz'), '4')
        border_el.set(qn('w:space'), '0')
        border_el.set(qn('w:color'), '000000')
        tcPr.append(border_el)


@print_doc_bp.route('/print/<int:employee_id>')
def print_service_record(employee_id):
    emp = Employee.query.get_or_404(employee_id)
    records = ServiceRecord.query.filter_by(employee_id=employee_id).order_by(
        ServiceRecord.sort_order
    ).all()

    certifier_name = _get_setting('certifier_name', 'PET ROANA B. BATACANDOLO, DPA')
    certifier_title = _get_setting('certifier_title', 'Designated HRMP')
    university_name = _get_setting('university_name', 'CAPIZ STATE UNIVERSITY')
    university_location = _get_setting('university_location', 'Pontevedra, Capiz')
    date_issued = datetime.now().strftime('%B %d, %Y')

    template_path = os.path.join(
        os.path.dirname(os.path.dirname(__file__)),
        'docx_templates',
        'service_record_template.docx',
    )
    if not os.path.exists(template_path):
        flash('DOCX template not found. Please run create_template.py first.', 'danger')
        return redirect(url_for('employees.view_employee', employee_id=employee_id))

    doc = Document(template_path)

    replacements = {
        'surname': emp.surname or '',
        'given_name': emp.given_name or '',
        'middle_name': emp.middle_name or '',
        'maiden_name': emp.maiden_name or '',
        'birth_date': emp.birth_date or '',
        'birth_place': emp.birth_place or '',
        'certifier_name': certifier_name,
        'certifier_title': certifier_title,
        'university_name': university_name,
        'university_location': university_location,
        'date_issued': date_issued,
    }

    # Replace in paragraphs
    for para in doc.paragraphs:
        _replace_text_in_paragraph(para, replacements)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_text_in_paragraph(para, replacements)

    # Find the service records table (the one with the marker row)
    service_table = None
    marker_row_idx = None
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            row_text = ' '.join(c.text for c in row.cells)
            if '{{service_records_table}}' in row_text:
                service_table = table
                marker_row_idx = i
                break
        if service_table is not None:
            break

    if service_table is not None and marker_row_idx is not None:
        # Remove the marker row
        marker_row_el = service_table.rows[marker_row_idx]._tr
        service_table._tbl.remove(marker_row_el)

        # Get the template data row (the row before marker, used as style reference)
        # We'll add new rows directly
        for rec in records:
            row = service_table.add_row()
            cells = row.cells
            data = [
                rec.date_from or '',
                rec.date_to or '',
                rec.designation or '',
                rec.status or '',
                _fmt_salary(rec.monthly_salary),
                _fmt_salary(rec.annual_salary),
                rec.station_place_of or '',
                rec.branch or '',
                rec.lv_abs_wo_pay or '',
                rec.separation_date or '',
                rec.separation_cause or '',
            ]
            for i, val in enumerate(data):
                if i < len(cells):
                    _set_cell_text(cells[i], val, font_size=8)
                    _set_cell_borders(cells[i])

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    filename = f'ServiceRecord_{emp.surname}_{emp.given_name}.docx'.replace(' ', '_')
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename,
    )
