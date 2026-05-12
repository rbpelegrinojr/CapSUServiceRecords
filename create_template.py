"""
create_template.py
Programmatically generates the service record DOCX template.
Run directly: python create_template.py
Or called by app.py on first launch.
"""
import os

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_para_format(para, bold=False, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                     space_before=0, space_after=0):
    para.alignment = alignment
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    if para.runs:
        run = para.runs[0]
        run.bold = bold
        run.font.size = Pt(size)


def _add_para(doc, text, bold=False, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER,
               space_before=0, space_after=0):
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return para


def _set_cell_text(cell, text, bold=False, size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                   vertical=WD_ALIGN_VERTICAL.CENTER):
    cell.text = ''
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)
    para = cell.paragraphs[0]
    para.alignment = alignment
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)


def _set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _merge_cells_horizontal(row, start, end):
    """Merge cells from start to end (inclusive) in a row."""
    start_cell = row.cells[start]
    end_cell = row.cells[end]
    start_cell.merge(end_cell)
    return start_cell


def create_template(output_path):
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(13)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # Remove default paragraph spacing
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    # ─── HEADER ────────────────────────────────────────────────────────────
    _add_para(doc, 'Republic of the Philippines', bold=False, size=11)
    _add_para(doc, '{{university_name}}', bold=True, size=13)
    _add_para(doc, '{{university_location}}', bold=False, size=11)
    _add_para(doc, '', size=6)
    _add_para(doc, 'SERVICE RECORDS', bold=True, size=14, space_after=6)

    # ─── EMPLOYEE INFO TABLE ────────────────────────────────────────────────
    emp_table = doc.add_table(rows=3, cols=4)
    emp_table.style = 'Table Grid'
    emp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(emp_table)

    # Row 0: Surname, Given Name
    r0 = emp_table.rows[0]
    _set_cell_text(r0.cells[0], 'SURNAME', bold=True, size=8)
    _set_cell_text(r0.cells[1], '{{surname}}', size=9,
                   alignment=WD_ALIGN_PARAGRAPH.LEFT)
    _set_cell_text(r0.cells[2], 'GIVEN NAME', bold=True, size=8)
    _set_cell_text(r0.cells[3], '{{given_name}}', size=9,
                   alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # Row 1: Middle Name, Maiden Name
    r1 = emp_table.rows[1]
    _set_cell_text(r1.cells[0], 'MIDDLE NAME', bold=True, size=8)
    _set_cell_text(r1.cells[1], '{{middle_name}}', size=9,
                   alignment=WD_ALIGN_PARAGRAPH.LEFT)
    _set_cell_text(r1.cells[2], 'MAIDEN NAME', bold=True, size=8)
    _set_cell_text(r1.cells[3], '{{maiden_name}}', size=9,
                   alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # Row 2: Birth Date, Birth Place
    r2 = emp_table.rows[2]
    _set_cell_text(r2.cells[0], 'DATE OF BIRTH', bold=True, size=8)
    _set_cell_text(r2.cells[1], '{{birth_date}}', size=9,
                   alignment=WD_ALIGN_PARAGRAPH.LEFT)
    _set_cell_text(r2.cells[2], 'PLACE OF BIRTH', bold=True, size=8)
    _set_cell_text(r2.cells[3], '{{birth_place}}', size=9,
                   alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # Spacer
    _add_para(doc, '', size=4)

    # ─── SERVICE RECORDS TABLE ──────────────────────────────────────────────
    # Columns: From | To | Designation | Status | Monthly | Annual | Station | Branch | LV ABS | Sep Date | Sep Cause
    # 11 columns total
    num_cols = 11
    sr_table = doc.add_table(rows=3, cols=num_cols)
    sr_table.style = 'Table Grid'
    sr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(sr_table)

    # Set column widths (approximate, total ~6.75 inches)
    col_widths = [Cm(2.0), Cm(2.0), Cm(3.2), Cm(1.8), Cm(1.8), Cm(1.8),
                  Cm(3.2), Cm(1.8), Cm(2.0), Cm(1.8), Cm(2.5)]

    for row in sr_table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(col_widths):
                cell.width = col_widths[i]

    # Row 0: Group headers
    r0 = sr_table.rows[0]
    # SERVICE spans col 0-1
    service_cell = _merge_cells_horizontal(r0, 0, 1)
    _set_cell_text(service_cell, 'SERVICE\n(Inclusive Dates)', bold=True, size=7)
    # RECORD OF APPOINTMENT spans col 2-5
    appt_cell = _merge_cells_horizontal(r0, 2, 5)
    _set_cell_text(appt_cell, 'RECORD OF APPOINTMENT', bold=True, size=7)
    # OFFICE ENTITY/DIVISION spans col 6-7
    office_cell = _merge_cells_horizontal(r0, 6, 7)
    _set_cell_text(office_cell, 'OFFICE\nENTITY/DIVISION', bold=True, size=7)
    # L/V ABS W/O PAY col 8
    _set_cell_text(r0.cells[8], 'L/V ABS\nW/O PAY', bold=True, size=7)
    # SEPARATION spans col 9-10
    sep_cell = _merge_cells_horizontal(r0, 9, 10)
    _set_cell_text(sep_cell, 'SEPARATION', bold=True, size=7)

    # Row 1: Sub-headers
    r1 = sr_table.rows[1]
    sub_headers = [
        'FROM', 'TO', 'DESIGNATION', 'STATUS',
        'MONTHLY\nSALARY', 'ANNUAL\nSALARY',
        'STATION/\nPLACE OF', 'BRANCH',
        '', 'DATE', 'CAUSE'
    ]
    for i, hdr in enumerate(sub_headers):
        _set_cell_text(r1.cells[i], hdr, bold=True, size=7)

    # Row 2: Marker row for data insertion
    r2 = sr_table.rows[2]
    marker_cell = _merge_cells_horizontal(r2, 0, num_cols - 1)
    _set_cell_text(marker_cell, '{{service_records_table}}', size=7)

    # ─── FOOTER ─────────────────────────────────────────────────────────────
    _add_para(doc, '', size=6)
    _add_para(
        doc,
        'Issued in compliance with Executive Order No. 54 dated August 10, 1954 '
        'and in conformity with the Memorandum of the Civil Service Commission '
        'dated January 8, 1985.',
        bold=False,
        size=8,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=4,
        space_after=4,
    )

    _add_para(doc, 'CERTIFIED CORRECT:', bold=True, size=10,
               alignment=WD_ALIGN_PARAGRAPH.LEFT)
    _add_para(doc, '', size=10)
    _add_para(doc, '{{certifier_name}}', bold=True, size=10,
               alignment=WD_ALIGN_PARAGRAPH.LEFT)
    _add_para(doc, '{{certifier_title}}', bold=False, size=10,
               alignment=WD_ALIGN_PARAGRAPH.LEFT)
    _add_para(doc, '', size=6)
    _add_para(doc, 'Date Issued: {{date_issued}}', bold=False, size=10,
               alignment=WD_ALIGN_PARAGRAPH.LEFT)

    doc.save(output_path)
    print(f'DOCX template created at: {output_path}')


if __name__ == '__main__':
    out = os.path.join(os.path.dirname(__file__), 'docx_templates', 'service_record_template.docx')
    create_template(out)
